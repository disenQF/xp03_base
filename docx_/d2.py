from docx import Document
import asyncio
import os


@asyncio.coroutine
def list_py(dirname):
    for filename in os.listdir(dirname):
        path = os.path.join(dirname, filename)
        if os.path.isdir(path):
            list_py(path)
        elif '.py' in os.path.splitext(path):
            yield from write_code(path)

    save('py_code.docx')


@asyncio.coroutine
def write_code(filename):
    doc.add_paragraph('源码文件-'+os.path.split(filename)[-1], style='List Number')
    tb = doc.add_table(rows=1, cols=1, style='Medium Shading 1 Accent 1')
    tb.rows[0].cells[0].text = '源码内容'
    with open(filename, 'rb') as code_file:
        code_text = code_file.read().decode()

        tb.add_row().cells[0].text = code_text

    p = doc.add_paragraph('\n')
    print('writing py', filename)


def save(filename):
    doc.save(filename)


if __name__ == '__main__':
    doc = Document()

    doc.add_heading('Python源码文档', level=0)

    loop = asyncio.get_event_loop()
    loop.run_until_complete(asyncio.wait((list_py('../process_'), )))
    loop.close()