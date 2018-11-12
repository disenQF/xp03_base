from docx import Document
import asyncio
import os


@asyncio.coroutine
def list_py(dirname):
    for filename in os.listdir(dirname):

        path = os.path.join(dirname, filename)
        print(path)
        if os.path.isdir(path):
            yield from list_py(path)
        elif '.java' in os.path.splitext(path):
            yield from write_code(path)

    save('nxu_pm_codes.docx')


@asyncio.coroutine
def write_code(filename):
    doc.add_paragraph('源码文件-'+os.path.split(filename)[-1], style='List Number')
    tb = doc.add_table(rows=1, cols=1, style='Medium Shading 1 Accent 1')
    tb.rows[0].cells[0].text = '源码内容'
    with open(filename, 'rb') as code_file:
        code_text = code_file.read().decode()

        tb.add_row().cells[0].text = code_text

    p = doc.add_paragraph('\n')
    print('writing ', filename)


def save(filename):
    doc.save(filename)


if __name__ == '__main__':
    doc = Document()

    doc.add_heading('科研项目源码文档', level=0)
    # pm_dir = '../process_'
    pm_dir = '/Users/apple/Documents/workspace-sts_1704/nxu_pm/src/main'

    loop = asyncio.get_event_loop()
    loop.run_until_complete(asyncio.wait((list_py(pm_dir), )))
    loop.close()