from docx import Document
from docx.shared import Pt, Inches


def save(doc, filename):
    print(type(doc))
    doc.save(filename)


def create():
    doc = Document()  # pip install python-docx

    doc.add_heading('源码文档', level=0)

    p = doc.add_paragraph('我的文件: aa.py')
    p.add_run('bold').bold = True
    p.add_run(' ----   ')
    p.add_run('我是斜体.').italic = True

    doc.add_heading('我的图片', level=1)
    doc.add_paragraph('Intense Quote', style='Intense Quote')

    doc.add_paragraph('第一个 未排序', style='List Bullet')
    doc.add_paragraph('第一个 排序', style='List Number')
    doc.add_picture('../process_/files/1.jpg', width=Pt(120))

    records = (
        (3, '101', 'Guido van Rossum'),
        (4, '102', 'Linus benedict Torvalds'),
        (5, '103', 'Bill Gates'),
    )

    tbl = doc.add_table(cols=3, rows=1, style='Medium Grid 1')  # 1行3列

    cells = tbl.rows[0].cells  # 获取1行的所有列（单元）
    cells[0].text = '序号'
    cells[1].text = '学号'
    cells[2].text = '姓名'

    for sn, id_ , name in records:
        r_cells = tbl.add_row().cells
        r_cells[0].text = str(sn)
        r_cells[1].text = id_
        r_cells[2].text = name

    doc.add_page_break()

    save(doc, 'py_code.docx')


if __name__ == '__main__':
    create()




